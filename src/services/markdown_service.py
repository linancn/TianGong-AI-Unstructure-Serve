"""Helpers for turning Markdown text into downloadable files."""

from __future__ import annotations

import os
import re
import subprocess
import tempfile
from functools import lru_cache
from pathlib import Path
from typing import Final

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

_DEFAULT_MARKDOWN_FILENAME: Final[str] = "document.md"
_DEFAULT_DOCX_FILENAME: Final[str] = "document.docx"
_GFM_PANDOC_FROM: Final[str] = (
    "gfm-hard_line_breaks+emoji"  # GitHub-flavoured markdown without hard breaks flattening headings
)
_HEADING_LINE_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
_TABLE_CAPTION_LEAD_PATTERN = re.compile(r"([^\n])\n(:[^\n]*\n)")
_TABLE_CAPTION_PATTERN = re.compile(r"(^|\n)(:[^\n]*?)\n(\|)")


def markdown_bytes(content: str, filename: str | None = None) -> tuple[str, bytes]:
    """Return a safe file name and UTF-8 bytes for the given Markdown content.

    If *filename* is provided the ``.md`` suffix is enforced and only the base
    name is kept so callers cannot escape the download directory.
    """

    if content is None:
        raise ValueError("Markdown content is required")

    safe_name = _safe_filename(filename, _DEFAULT_MARKDOWN_FILENAME, ".md")

    return safe_name, content.encode("utf-8")


def markdown_to_docx_bytes(
    content: str,
    filename: str | None = None,
    reference_doc_path: str | None = None,
) -> tuple[str, bytes]:
    """Return DOCX bytes converted from Markdown using Pandoc.

    Accepts an optional *reference_doc_path* to control styling and falls back
    to the bundled DOCX template when none is supplied.

    After conversion, post-processes the document to ensure styles are properly
    applied and direct formatting is removed for consistent Chinese-English text display.
    """

    if content is None:
        raise ValueError("Markdown content is required")

    safe_name = _safe_filename(filename, _DEFAULT_DOCX_FILENAME, ".docx")
    normalized_content = _normalize_markdown(content)
    reference_doc = _resolve_reference_doc(reference_doc_path)

    pandoc_cmd = [_pandoc_executable(), f"--from={_pandoc_from()}", "--to=docx"]

    for lua_filter in _pandoc_filters():
        pandoc_cmd.extend(["--lua-filter", lua_filter])

    if reference_doc:
        pandoc_cmd.extend(["--reference-doc", reference_doc])

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        output_path = tmpdir_path / "output.docx"
        pandoc_cmd.extend(["--output", str(output_path)])

        try:
            subprocess.run(
                pandoc_cmd,
                input=normalized_content,
                text=True,
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
        except FileNotFoundError as exc:  # Pandoc missing from environment.
            raise RuntimeError(
                "Pandoc executable not found. Install pandoc or set PANDOC_PATH"
            ) from exc
        except subprocess.CalledProcessError as exc:
            raise RuntimeError(
                "Pandoc failed to create DOCX: {error}".format(error=exc.stderr)
            ) from exc

        # Post-process the document to fix style issues
        if DOCX_AVAILABLE:
            _fix_document_styles(str(output_path))

        docx_bytes = output_path.read_bytes()

    return safe_name, docx_bytes


def _safe_filename(filename: str | None, default: str, suffix: str) -> str:
    """Return an os-safe file name with the desired suffix enforced."""

    safe_name = (filename or default).strip() or default
    safe_name = os.path.basename(safe_name)

    if not safe_name.lower().endswith(suffix):
        safe_name = f"{safe_name}{suffix}"

    return safe_name


def _fix_document_styles(docx_path: str) -> None:
    """Post-process DOCX to ensure styles are properly applied without direct formatting.

    This fixes issues where:
    1. Styles need manual updating in Word to display properly
    2. Chinese-English mixed text doesn't display correctly
    3. Styles revert after saving and reopening

    The fix works by:
    - Removing direct formatting (rPr) from paragraph runs that conflicts with styles
    - Adding proper font definitions to runs for styles that don't define fonts
    - Ensuring paragraph styles are properly linked
    - Preserving only essential formatting like bold, italic, code
    """
    if not DOCX_AVAILABLE:
        return

    try:
        doc = Document(docx_path)

        # Process all paragraphs
        for paragraph in doc.paragraphs:
            _clean_and_enhance_paragraph_style(paragraph, doc)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _clean_and_enhance_paragraph_style(paragraph, doc)

        # Save the cleaned document
        doc.save(docx_path)
    except Exception as e:
        # If post-processing fails, log but don't break the conversion
        print(f"Warning: Could not post-process DOCX styles: {e}")


def _clean_and_enhance_paragraph_style(paragraph, doc) -> None:
    """Clean a paragraph and add proper font definitions if the style doesn't have them.

    Strategy:
    1. Remove problematic direct formatting that conflicts with styles
    2. For styles without font definitions, add proper Chinese-English fonts to runs
    3. Preserve semantic formatting (bold, italic, underline)
    """

    # Get the paragraph's style name
    style_name = paragraph.style.name if paragraph.style else None

    # Check if this style has font definitions
    style_has_fonts = False
    if style_name:
        try:
            style = doc.styles[style_name]
            style_element = style._element
            rPr = style_element.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    # Check if it has any font defined
                    if (
                        rFonts.get(qn("w:ascii"))
                        or rFonts.get(qn("w:eastAsia"))
                        or rFonts.get(qn("w:hAnsi"))
                    ):
                        style_has_fonts = True
        except:
            pass

    # Elements to remove (these often conflict with style definitions)
    ELEMENTS_TO_REMOVE = {
        qn("w:sz"),  # Font size
        qn("w:szCs"),  # Font size for complex scripts
        qn("w:color"),  # Font color
        qn("w:spacing"),  # Character spacing
        qn("w:kern"),  # Kerning
        qn("w:w"),  # Character width scaling
        qn("w:position"),  # Position
    }

    # For each run in the paragraph
    for run in paragraph.runs:
        rPr = run._element.rPr

        # If the style doesn't have fonts, we need to add them to the run
        if not style_has_fonts:
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                run._element.insert(0, rPr)

            # Remove existing rFonts if any (to start fresh)
            existing_rFonts = rPr.find(qn("w:rFonts"))
            if existing_rFonts is not None:
                rPr.remove(existing_rFonts)

            # Add proper font definitions for Chinese-English mixed text
            rFonts = OxmlElement("w:rFonts")
            # Use Times New Roman for English and SimSun (宋体) for Chinese
            rFonts.set(qn("w:ascii"), "Times New Roman")  # English font
            rFonts.set(qn("w:eastAsia"), "宋体")  # Chinese font
            rFonts.set(qn("w:hAnsi"), "Times New Roman")  # High ANSI font
            rFonts.set(qn("w:cs"), "Times New Roman")  # Complex scripts font

            # Add Word-specific hint attribute for better compatibility
            # This tells Word to prefer East Asian fonts for mixed content
            rFonts.set(qn("w:hint"), "eastAsia")

            # Insert rFonts as the first element in rPr
            rPr.insert(0, rFonts)

        # Clean up problematic formatting elements
        if rPr is not None:
            # Remove problematic elements
            for element in list(rPr):  # Use list() to avoid modifying during iteration
                if element.tag in ELEMENTS_TO_REMOVE:
                    rPr.remove(element)

    # Ensure the paragraph style is properly set
    if style_name:
        try:
            paragraph.style = style_name
        except:
            pass  # Style might not exist, keep current


def _resolve_reference_doc(reference_doc_path: str | None) -> str | None:
    """Prefer explicit reference doc path, fall back to packaged default template."""

    candidate_path = reference_doc_path

    if not candidate_path:
        default_candidate = Path(__file__).resolve().parent / "templates" / "default_reference.docx"
        return str(default_candidate) if default_candidate.exists() else None

    resolved_path = Path(candidate_path).expanduser().resolve()
    if not resolved_path.exists():
        raise ValueError(f"Reference DOCX template not found: {resolved_path}")

    return str(resolved_path)


def _pandoc_filters() -> list[str]:
    filters: list[str] = []
    filters_dir = Path(__file__).resolve().parent / "filters"

    for name in ("pagebreak.lua", "figure_caption.lua"):
        candidate = filters_dir / name
        if candidate.exists():
            filters.append(str(candidate))

    return filters


def _pandoc_executable() -> str:
    """Allow overriding the pandoc executable via environment variable."""

    return os.getenv("PANDOC_PATH", "pandoc")


@lru_cache(maxsize=1)
def _pandoc_from() -> str:
    """Choose a Pandoc input format; stay on GFM, add raw_tex only if supported."""

    raw_tex = "raw_tex"

    if _pandoc_supports_extension("gfm", raw_tex):
        return f"{_GFM_PANDOC_FROM}+{raw_tex}"

    return _GFM_PANDOC_FROM


def _pandoc_supports_extension(format_name: str, extension: str) -> bool:
    try:
        proc = subprocess.run(
            [_pandoc_executable(), f"--list-extensions={format_name}"],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except (FileNotFoundError, subprocess.SubprocessError):
        return False

    for line in proc.stdout.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        if stripped[0] in {"+", "-"}:
            stripped = stripped[1:]

        if stripped == extension:
            return True

    return False


def _normalize_markdown(content: str) -> str:
    """Tidy headings so Pandoc reliably recognises heading levels."""

    content = _TABLE_CAPTION_LEAD_PATTERN.sub(
        lambda match: f"{match.group(1)}\n\n{match.group(2)}", content
    )
    content = _TABLE_CAPTION_PATTERN.sub(
        lambda match: f"{match.group(1)}{match.group(2)}\n\n{match.group(3)}", content
    )

    lines = content.splitlines()
    normalised_lines: list[str] = []
    last_was_heading = False

    for raw_line in lines:
        line = raw_line.rstrip()
        heading_match = _HEADING_LINE_PATTERN.match(line)

        if heading_match:
            if normalised_lines and normalised_lines[-1]:
                normalised_lines.append("")
            normalised_lines.append(f"{heading_match.group(1)} {heading_match.group(2).strip()}")
            last_was_heading = True
            continue

        if line:
            if last_was_heading and normalised_lines and normalised_lines[-1]:
                normalised_lines.append("")
            normalised_lines.append(line)
            last_was_heading = False
        else:
            if not normalised_lines or normalised_lines[-1]:
                normalised_lines.append("")
            last_was_heading = False

    return "\n".join(normalised_lines).strip("\n") + "\n"
